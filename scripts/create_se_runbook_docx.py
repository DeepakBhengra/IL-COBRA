#!/usr/bin/env python3
"""Create samples/docs/runbook-SE-terms-override.docx for SE ingest testing."""

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "samples" / "docs" / "runbook-SE-terms-override.docx"


def main() -> None:
    doc = Document()
    doc.add_heading("Runbook: CORORA Error SE — Terms Security Override", 0)
    doc.add_paragraph(
        "Program ORP676 sets CORORA-R-ERR-NO-SEC-TERM-OVRD when the operator lacks "
        "authority to override payment terms. Two-character error code: SE."
    )
    doc.add_heading("When you see error code SE", level=1)
    for text in (
        "Confirm program ORP676 and transaction TERM (TB-SEC-FIELD-NBR 05).",
        "Verify 88-level CORORA-R-ERR-NO-SEC-TERM-OVRD / alias ERR-NO-SEC-TERM-OVRD.",
        "Check operator profile for SEC-TERM-05 grant in security admin.",
        "After grant, retest terms change; SE should not recur.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Mapping reference", level=1)
    doc.add_paragraph(
        "CORORA_TWO_CHAR_ERROR.txt: VALUE 'SE' -> CORORA-R-ERR-NO-SEC-TERM-OVRD"
    )
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
